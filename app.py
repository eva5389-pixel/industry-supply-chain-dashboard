import re
from datetime import datetime
from io import BytesIO
import pandas as pd
import streamlit as st
import yfinance as yf
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# 設定網頁版面與暗色風格
st.set_page_config(
    page_title="全球科技與供應鏈即時漲跌儀表板", page_icon="📈", layout="wide"
)

# 注入自定義 CSS
st.markdown(
    """
    <style>
    .main { background-color: #0e1117; color: #ffffff; }
    .stMetric { background-color: #161b22; padding: 15px; border-radius: 8px; border: 1px solid #30363d; }
    </style>
""",
    unsafe_allow_html=True,
)

st.title("🌐 全球科技、機器人、軍工、低軌衛星與供應鏈即時漲跌總覽")
st.markdown(
    "即時串接 Yahoo Finance，追蹤美、台、日、中、港各關鍵產業供應鏈、上下游定位與承作業務。"
)

# 定義跨國供應鏈清單（移除所有名稱後面的括號說明）
supply_chains = {
    "量子電腦": [
        {
            "名稱": "台積電",
            "位置": "下游",
            "業務": "量子晶片代工與先進製程",
            "代碼": "2330.TW",
        },
        {
            "名稱": "日月光投控",
            "位置": "下游",
            "業務": "量子晶片先進封裝與測試",
            "代碼": "3711.TW",
        },
        {
            "名稱": "千附精密",
            "位置": "上游",
            "業務": "超低溫量子設備系統製造",
            "代碼": "6829.TW",
        },
        {
            "名稱": "金寶",
            "位置": "中游",
            "業務": "SEEQC投資/控制晶片開發",
            "代碼": "2312.TW",
        },
        {
            "名稱": "康舒",
            "位置": "中游",
            "業務": "室溫與低溫控制電源裝置",
            "代碼": "6282.TW",
        },
        {
            "名稱": "IBM",
            "位置": "下游",
            "業務": "量子硬體與雲端運算服務",
            "代碼": "IBM",
        },
        {
            "名稱": "Google/Alphabet",
            "位置": "下游",
            "業務": "超導量子處理器研發",
            "代碼": "GOOGL",
        },
        {
            "名稱": "IonQ",
            "位置": "中游",
            "業務": "離子阱量子計算系統",
            "代碼": "IONQ",
        },
        {
            "名稱": "Rigetti",
            "位置": "中游",
            "業務": "超導量子計算機",
            "代碼": "RGTI",
        },
        {
            "名稱": "D-Wave",
            "位置": "中游",
            "業務": "量子退火系統",
            "代碼": "QBTS",
        },
        {
            "名稱": "廣達",
            "位置": "下游",
            "業務": "系統整合與國際大廠結盟",
            "代碼": "2382.TW",
        },
        {
            "名稱": "鴻海",
            "位置": "下游",
            "業務": "系統整合與量子運算應用",
            "代碼": "2317.TW",
        },
    ],
    "低軌衛星": [
        {
            "名稱": "SpaceX",
            "位置": "下游",
            "業務": "星鏈衛星網路運營/特斯拉概念",
            "代碼": "TSLA",
        },
        {
            "名稱": "思佳訊",
            "位置": "上游",
            "業務": "射頻晶片與衛星通訊零組件",
            "代碼": "SWKS",
        },
        {
            "名稱": "Viasat",
            "位置": "下游",
            "業務": "衛星寬頻與國防通訊服務",
            "代碼": "VSAT",
        },
        {
            "名稱": "升達科",
            "位置": "上游",
            "業務": "衛星通訊高頻被動元件/濾波器",
            "代碼": "3491.TWO",
        },
        {
            "名稱": "耀登",
            "位置": "上游",
            "業務": "衛星陣列天線與通訊測試",
            "代碼": "3138.TW",
        },
        {
            "名稱": "芳興",
            "位置": "上游",
            "業務": "衛星結構件與精密機械加工",
            "代碼": "4526.TW",
        },
        {
            "名稱": "NEC",
            "位置": "下游",
            "業務": "日本衛星系統與通訊設備整合",
            "代碼": "6701.T",
        },
    ],
    "軍工/國防": [
        {
            "名稱": "洛歇馬丁",
            "位置": "下游",
            "業務": "匿蹤戰機與導彈系統",
            "代碼": "LMT",
        },
        {
            "名稱": "雷神技術",
            "位置": "下游",
            "業務": "防空飛彈與航太防衛",
            "代碼": "RTX",
        },
        {
            "名稱": "諾斯洛普格魯曼",
            "位置": "下游",
            "業務": "轟炸機與無人機系統",
            "代碼": "NOC",
        },
        {
            "名稱": "通用動力",
            "位置": "下游",
            "業務": "裝甲車、潛艦與戰車",
            "代碼": "GD",
        },
        {
            "名稱": "帕蘭提爾",
            "位置": "中游",
            "業務": "國防AI大數據分析軟體",
            "代碼": "PLTR",
        },
        {
            "名稱": "漢翔",
            "位置": "中游",
            "業務": "軍用機維修與機體結構",
            "代碼": "2634.TW",
        },
        {
            "名稱": "雷虎",
            "位置": "中游",
            "業務": "軍用無人機與無人艇研發",
            "代碼": "8033.TW",
        },
        {
            "名稱": "駐龍",
            "位置": "上游",
            "業務": "航太結構件與機身零組件",
            "代碼": "4572.TW",
        },
        {
            "名稱": "寶一",
            "位置": "上游",
            "業務": "航太引擎零件與發動機金屬件",
            "代碼": "8222.TW",
        },
        {
            "名稱": "三菱重工業",
            "位置": "下游",
            "業務": "日本防衛裝備與艦艇",
            "代碼": "7011.T",
        },
        {
            "名稱": "川崎重工業",
            "位置": "下游",
            "業務": "日本軍用直升機與潛艦",
            "代碼": "7012.T",
        },
        {
            "名稱": "萊茵金屬",
            "位置": "下游",
            "業務": "歐洲軍火與重型武器",
            "代碼": "RHM.DE",
        },
        {
            "名稱": "航發動力",
            "位置": "上游",
            "業務": "中國航空發動機與零組件",
            "代碼": "600893.SS",
        },
    ],
    "半導體": [
        {
            "名稱": "台積電",
            "位置": "中游",
            "業務": "全球晶圓代工與先進製程",
            "代碼": "2330.TW",
        },
        {
            "名稱": "力積電",
            "位置": "中游",
            "業務": "晶圓代工與記憶體製造",
            "代碼": "6770.TW",
        },
        {
            "名稱": "大立光",
            "位置": "下游",
            "業務": "手機鏡頭與光學元件",
            "代碼": "3008.TW",
        },
        {
            "名稱": "日電貿",
            "位置": "上游",
            "業務": "被動元件與電容代理經銷",
            "代碼": "3090.TW",
        },
        {
            "名稱": "NVIDIA",
            "位置": "下游",
            "業務": "AI繪圖晶片與高效運算處理器",
            "代碼": "NVDA",
        },
        {
            "名稱": "艾司摩爾",
            "位置": "上游",
            "業務": "極紫外光EUV微影曝光機",
            "代碼": "ASML",
        },
        {
            "名稱": "艾德萬測試",
            "位置": "上游",
            "業務": "半導體自動測試設備ATE",
            "代碼": "6857.T",
        },
        {
            "名稱": "中芯國際",
            "位置": "中游",
            "業務": "中國晶圓代工製造",
            "代碼": "0981.HK",
        },
        {
            "名稱": "東京威力科創",
            "位置": "上游",
            "業務": "半導體蝕刻與成膜設備",
            "代碼": "8035.T",
        },
    ],
    "晶片設計": [
        {
            "名稱": "聯發科",
            "位置": "上游",
            "業務": "行動處理器與IC設計",
            "代碼": "2454.TW",
        },
        {
            "名稱": "創意",
            "位置": "上游",
            "業務": "特殊應用晶片ASIC與IP設計",
            "代碼": "3443.TW",
        },
    ],
    "AI伺服器/雲端": [
        {"名稱": "仁寶", "位置": "中游", "業務": "伺服器、筆電與電子產品ODM製造", "代碼": "2324.TW"},
        {"名稱": "川湖", "位置": "上游", "業務": "AI伺服器滑軌、機構件與導軌系統", "代碼": "2059.TW"},
        {"名稱": "CoreWeave", "位置": "下游", "業務": "GPU算力租賃、AI雲端基礎設施與訓練平台", "代碼": "CRWV"},
        {"名稱": "Palantir", "位置": "下游", "業務": "企業與國防AI資料分析、決策軟體平台", "代碼": "PLTR"},
    ],
    "光通訊": [
        {"名稱": "聯亞", "位置": "上游", "業務": "InP磊晶片、雷射元件與高速光通訊材料", "代碼": "3081.TWO"},
        {"名稱": "全新", "位置": "上游", "業務": "砷化鎵與磷化銦磊晶片、光電子材料", "代碼": "2455.TW"},
        {"名稱": "環宇-KY", "位置": "上游", "業務": "砷化鎵與磷化銦晶圓、光通訊元件代工", "代碼": "4991.TW"},
        {"名稱": "Lumentum", "位置": "上游", "業務": "InP雷射晶片、高速光模組與光通訊元件", "代碼": "LITE"},
        {"名稱": "Coherent", "位置": "上游", "業務": "磷化銦材料、雷射光源與光通訊平台", "代碼": "COHR"},
        {"名稱": "源杰科技", "位置": "上游", "業務": "高速光半導體雷射器晶片與EML/CW光源", "代碼": "688498.SS"},
        {"名稱": "光庫科技", "位置": "上游", "業務": "光纖元件、鈮酸鋰調製器與光學晶片", "代碼": "300620.SZ"},
        {"名稱": "光環", "位置": "中游", "業務": "高速光模組、雷射二極體與光通訊元件", "代碼": "3234.TWO"},
        {"名稱": "華星光", "位置": "中游", "業務": "高速光收發模組與光通訊次模組", "代碼": "4979.TWO"},
        {"名稱": "上詮", "位置": "中游", "業務": "光纖被動元件、連接器與矽光子耦合元件", "代碼": "3363.TWO"},
        {"名稱": "光聖", "位置": "中游", "業務": "光纖連接器、資料中心與高階光通訊產品", "代碼": "6442.TW"},
        {"名稱": "波若威", "位置": "中游", "業務": "光纖被動元件、分光器與資料中心光通訊", "代碼": "3163.TWO"},
        {"名稱": "眾達-KY", "位置": "中游", "業務": "高速光收發模組與資料中心光通訊產品", "代碼": "4977.TW"},
        {"名稱": "Applied Optoelectronics", "位置": "中游", "業務": "垂直整合高速光模組、雷射與資料中心產品", "代碼": "AAOI"},
        {"名稱": "Marvell", "位置": "中游", "業務": "網通晶片、DSP與CPO光引擎設計", "代碼": "MRVL"},
        {"名稱": "Broadcom", "位置": "中游", "業務": "交換器晶片、網通ASIC與CPO光互連平台", "代碼": "AVGO"},
        {"名稱": "中際旭創", "位置": "中游", "業務": "800G／1.6T高速光模組與資料中心互連", "代碼": "300308.SZ"},
        {"名稱": "新易盛", "位置": "中游", "業務": "高速光模組與海外資料中心光互連產品", "代碼": "300502.SZ"},
        {"名稱": "天孚通信", "位置": "中游", "業務": "光引擎、無源光器件與精密光學元件", "代碼": "300394.SZ"},
        {"名稱": "長飛光纖", "位置": "中游", "業務": "光纖預製棒、光纖與光纜製造", "代碼": "601869.SS"},
        {"名稱": "亨通光電", "位置": "中游", "業務": "光纖光纜、海纜與通信網路產品", "代碼": "600487.SS"},
        {"名稱": "藤倉", "位置": "中游", "業務": "光纖、電纜與資料中心高速互連產品", "代碼": "5803.T"},
        {"名稱": "古河電工", "位置": "中游", "業務": "光纖網路、光纜與資料中心線材", "代碼": "5801.T"},
        {"名稱": "住友電工", "位置": "中游", "業務": "光纖通訊、汽車線束與電力設備", "代碼": "5802.T"},
        {"名稱": "日月光投控", "位置": "中下游", "業務": "光通訊晶片封裝、測試與先進封裝", "代碼": "3711.TW"},
        {"名稱": "訊芯-KY", "位置": "中下游", "業務": "高速光通訊與半導體系統級封裝", "代碼": "6451.TW"},
        {"名稱": "聯鈞", "位置": "中下游", "業務": "光通訊元件封裝、雷射模組與測試", "代碼": "3450.TW"},
        {"名稱": "旺矽", "位置": "中下游", "業務": "半導體與光通訊測試介面、探針卡", "代碼": "6223.TW"},
    ],
    "顯示面板": [
        {"名稱": "群創", "位置": "中游", "業務": "顯示面板、車用顯示與先進封裝相關應用", "代碼": "3481.TW"},
    ],
    "ASIC設計": [
        {"名稱": "世芯-KY", "位置": "上游", "業務": "高階ASIC設計服務與先進製程SoC整合", "代碼": "3661.TW"},
    ],
    "日本五大商社": [
        {"名稱": "三菱商事", "位置": "跨供應鏈", "業務": "能源、金屬、食品、基礎建設與全球貿易投資", "代碼": "8058.T"},
        {"名稱": "三井物產", "位置": "跨供應鏈", "業務": "資源、能源、機械、化學與全球事業投資", "代碼": "8031.T"},
        {"名稱": "伊藤忠商事", "位置": "跨供應鏈", "業務": "消費、紡織、食品、機械與資訊通訊事業", "代碼": "8001.T"},
        {"名稱": "住友商事", "位置": "跨供應鏈", "業務": "金屬、運輸、基礎建設、媒體與不動產事業", "代碼": "8053.T"},
        {"名稱": "丸紅", "位置": "跨供應鏈", "業務": "電力、糧食、化學、金屬與全球貿易投資", "代碼": "8002.T"},
    ],
    "中國8Tech": [
        {"名稱": "騰訊", "位置": "下游", "業務": "社群、遊戲、雲端與金融科技平台", "代碼": "0700.HK"},
        {"名稱": "阿里巴巴", "位置": "下游", "業務": "電商、雲端運算、物流與AI平台", "代碼": "9988.HK"},
        {"名稱": "美團", "位置": "下游", "業務": "本地生活、外送、旅遊與即時零售平台", "代碼": "3690.HK"},
        {"名稱": "小米", "位置": "下游", "業務": "智慧手機、AIoT、電動車與消費電子", "代碼": "1810.HK"},
        {"名稱": "京東", "位置": "下游", "業務": "電商、物流、供應鏈與雲端科技", "代碼": "9618.HK"},
        {"名稱": "百度", "位置": "下游", "業務": "搜尋、生成式AI、智能雲與自動駕駛", "代碼": "9888.HK"},
        {"名稱": "網易", "位置": "下游", "業務": "線上遊戲、音樂、教育與網路服務", "代碼": "9999.HK"},
        {"名稱": "中芯國際", "位置": "中游", "業務": "中國晶圓代工與成熟製程製造", "代碼": "0981.HK"},
    ],
    "電源/封裝測試": [
        {"名稱": "台達電", "位置": "中游", "業務": "電源管理、資料中心電源、散熱與能源自動化方案", "代碼": "2308.TW"},
        {"名稱": "京元電子", "位置": "下游", "業務": "IC測試、晶圓測試與半導體後段測試服務", "代碼": "2449.TW"},
    ],
    "台灣金融": [
        {"名稱": "第一金", "位置": "下游", "業務": "銀行、證券、投信與綜合金融服務", "代碼": "2892.TW"},
        {"名稱": "合庫金", "位置": "下游", "業務": "商業銀行、保險、證券與資產管理", "代碼": "5880.TW"},
        {"名稱": "中信金", "位置": "下游", "業務": "銀行、信用卡、保險與財富管理", "代碼": "2891.TW"},
        {"名稱": "玉山金", "位置": "下游", "業務": "銀行、信用卡、數位金融與財富管理", "代碼": "2884.TW"},
    ],
    "醫療生技": [
        {"名稱": "藥華藥", "位置": "中游", "業務": "蛋白質新藥研發、製造與血液疾病藥物商業化", "代碼": "6446.TW"},
        {"名稱": "保瑞", "位置": "中游", "業務": "國際藥品CDMO、製劑生產與藥品銷售", "代碼": "6472.TWO"},
        {"名稱": "台康生技", "位置": "中游", "業務": "生物相似藥、抗體藥物與生物製劑CDMO", "代碼": "6589.TWO"},
        {"名稱": "合一", "位置": "上游", "業務": "新藥研發、傷口照護與慢性病藥物開發", "代碼": "4743.TWO"},
        {"名稱": "台耀", "位置": "上游", "業務": "原料藥、特殊學名藥與委託開發製造", "代碼": "4746.TWO"},
        {"名稱": "智擎", "位置": "上游", "業務": "癌症新藥研發、授權與國際臨床開發", "代碼": "4162.TWO"},
        {"名稱": "Eli Lilly", "位置": "下游", "業務": "糖尿病、肥胖症、腫瘤與神經疾病創新藥物", "代碼": "LLY"},
        {"名稱": "Novo Nordisk", "位置": "下游", "業務": "糖尿病、肥胖症與慢性疾病生物製劑", "代碼": "NVO"},
        {"名稱": "Johnson & Johnson", "位置": "下游", "業務": "創新藥物、醫療器材與全球醫療產品", "代碼": "JNJ"},
        {"名稱": "Merck", "位置": "下游", "業務": "腫瘤免疫、疫苗與處方藥研發銷售", "代碼": "MRK"},
        {"名稱": "Amgen", "位置": "中游", "業務": "生物製劑、抗體藥物與重大疾病新藥", "代碼": "AMGN"},
        {"名稱": "Moderna／莫德納", "位置": "中游", "業務": "mRNA疫苗、傳染病疫苗與個人化癌症治療研發", "代碼": "MRNA"},
        {"名稱": "UnitedHealth", "位置": "下游", "業務": "健康保險、醫療服務與醫療數據管理", "代碼": "UNH"},
    ],
    "記憶體": [
        {"名稱": "南亞科", "位置": "中游", "業務": "DRAM記憶體設計與製造", "記憶體類型": "DDR／DRAM", "代碼": "2408.TW"},
        {"名稱": "華邦電", "位置": "中游", "業務": "利基型DRAM與編碼型快閃記憶體", "記憶體類型": "DDR／NOR NAND", "代碼": "2344.TW"},
        {"名稱": "威剛", "位置": "下游", "業務": "記憶體模組、SSD與工業儲存產品", "記憶體類型": "DDR／NAND", "代碼": "3260.TWO"},
        {"名稱": "力成", "位置": "下游", "業務": "記憶體封裝測試與先進封裝服務", "記憶體類型": "NAND／DDR／HBM封測", "代碼": "6239.TW"},
        {"名稱": "Micron", "位置": "中游", "業務": "DRAM、NAND與高頻寬記憶體製造", "記憶體類型": "HBM／DDR／NAND", "代碼": "MU"},
        {"名稱": "SK海力士", "位置": "中游", "業務": "高頻寬記憶體、DRAM與NAND製造", "記憶體類型": "HBM／DDR／NAND", "代碼": "000660.KS"},
        {"名稱": "Samsung Electronics", "位置": "中游", "業務": "DRAM、HBM、NAND與消費電子", "記憶體類型": "HBM／DDR／NAND", "代碼": "005930.KS"},
        {"名稱": "SanDisk", "位置": "中游", "業務": "NAND快閃記憶體、SSD與儲存產品", "記憶體類型": "NAND", "代碼": "SNDK"},
    ],
    "輝達供應鏈": [
        {"名稱": "NVIDIA", "位置": "上游", "業務": "GPU、AI加速晶片與運算平台", "代碼": "NVDA"},
        {"名稱": "台積電", "位置": "中游", "業務": "GPU先進製程晶圓代工與CoWoS封裝", "代碼": "2330.TW"},
        {"名稱": "鴻海", "位置": "中游", "業務": "AI伺服器與整機櫃製造", "代碼": "2317.TW"},
        {"名稱": "廣達", "位置": "中游", "業務": "AI伺服器ODM與資料中心系統", "代碼": "2382.TW"},
        {"名稱": "緯創", "位置": "中游", "業務": "AI伺服器、加速卡與系統組裝", "代碼": "3231.TW"},
        {"名稱": "光寶科", "位置": "上游", "業務": "伺服器電源、散熱與電源管理模組", "代碼": "2301.TW"},
        {"名稱": "川湖", "位置": "上游", "業務": "AI伺服器滑軌與機構件", "代碼": "2059.TW"},
    ],
    "博通供應鏈": [
        {"名稱": "Broadcom", "位置": "上游", "業務": "網通晶片、交換器ASIC與客製化加速晶片", "代碼": "AVGO"},
        {"名稱": "台積電", "位置": "中游", "業務": "高階網通與ASIC晶圓代工", "代碼": "2330.TW"},
        {"名稱": "日月光投控", "位置": "下游", "業務": "網通晶片封裝測試與系統級封裝", "代碼": "3711.TW"},
        {"名稱": "智邦", "位置": "下游", "業務": "高速交換器與資料中心網路設備", "代碼": "2345.TW"},
        {"名稱": "啟碁", "位置": "中游", "業務": "無線通訊、網通模組與企業網路設備", "代碼": "6285.TW"},
        {"名稱": "南電", "位置": "上游", "業務": "ABF載板與高階IC封裝基板", "代碼": "8046.TW"},
        {"名稱": "Marvell", "位置": "上游", "業務": "資料中心互連、網通與客製化運算晶片", "代碼": "MRVL"},
    ],
    "SpaceX供應鏈": [
        {"名稱": "SpaceX概念／Tesla", "位置": "下游", "業務": "低軌衛星發射、Starlink網路與終端服務概念", "代碼": "TSLA"},
        {"名稱": "升達科", "位置": "上游", "業務": "衛星微波與毫米波通訊元件", "代碼": "3491.TW"},
        {"名稱": "耀登", "位置": "上游", "業務": "衛星天線、射頻測試與通訊元件", "代碼": "3138.TWO"},
        {"名稱": "Viasat", "位置": "下游", "業務": "衛星寬頻與國防通訊服務", "代碼": "VSAT"},
        {"名稱": "Skyworks", "位置": "上游", "業務": "射頻前端晶片與衛星通訊元件", "代碼": "SWKS"},
    ],
    "CoWoS供應鏈": [
        {"名稱": "台積電", "位置": "中游", "業務": "CoWoS先進封裝與先進製程晶圓製造", "代碼": "2330.TW"},
        {"名稱": "日月光投控", "位置": "下游", "業務": "先進封裝、測試與系統級封裝", "代碼": "3711.TW"},
        {"名稱": "京鼎", "位置": "上游", "業務": "半導體製程與先進封裝設備模組", "代碼": "3413.TW"},
        {"名稱": "辛耘", "位置": "上游", "業務": "濕製程、再生晶圓與先進封裝設備", "代碼": "3583.TW"},
        {"名稱": "萬潤", "位置": "上游", "業務": "半導體與先進封裝自動化設備", "代碼": "6187.TWO"},
        {"名稱": "弘塑", "位置": "上游", "業務": "先進封裝濕製程設備與化學供應系統", "代碼": "3131.TWO"},
        {"名稱": "景碩", "位置": "上游", "業務": "ABF載板與高階IC封裝基板", "代碼": "3189.TW"},
    ],
    "玻璃纖維": [
        {"名稱": "台玻", "位置": "上游", "業務": "電子級玻璃纖維布、玻璃材料與建築玻璃", "代碼": "1802.TW"},
        {"名稱": "富喬", "位置": "上游", "業務": "電子級玻璃纖維紗與玻纖布", "代碼": "1815.TWO"},
    ],
    "英特爾供應鏈": [
        {"名稱": "Intel", "位置": "中游", "業務": "CPU設計、晶圓製造與先進封裝", "代碼": "INTC"},
        {"名稱": "ASML", "位置": "上游", "業務": "EUV與DUV微影設備", "代碼": "ASML"},
        {"名稱": "Applied Materials", "位置": "上游", "業務": "沉積、蝕刻與半導體製程設備", "代碼": "AMAT"},
        {"名稱": "Lam Research", "位置": "上游", "業務": "蝕刻、薄膜與晶圓清洗設備", "代碼": "LRCX"},
        {"名稱": "台積電", "位置": "中游", "業務": "部分處理器與晶片晶圓代工", "代碼": "2330.TW"},
        {"名稱": "日月光投控", "位置": "下游", "業務": "封裝測試與系統級封裝服務", "代碼": "3711.TW"},
    ],
    "成熟製程": [
        {"名稱": "聯電", "位置": "中游", "業務": "成熟製程與特殊製程晶圓代工", "代碼": "2303.TW"},
        {"名稱": "世界先進", "位置": "中游", "業務": "成熟製程、電源管理與特殊製程晶圓代工", "代碼": "5347.TWO"},
        {"名稱": "力積電", "位置": "中游", "業務": "成熟製程、記憶體與邏輯晶圓代工", "代碼": "6770.TW"},
        {"名稱": "中芯國際", "位置": "中游", "業務": "中國成熟及先進節點晶圓代工", "代碼": "0981.HK"},
        {"名稱": "華虹半導體", "位置": "中游", "業務": "功率、嵌入式記憶體與成熟製程代工", "代碼": "1347.HK"},
        {"名稱": "穩懋", "位置": "中游", "業務": "砷化鎵與射頻晶圓代工", "代碼": "3105.TWO"},
    ],
    "先進製程": [
        {"名稱": "台積電", "位置": "中游", "業務": "2奈米、3奈米及先進製程晶圓代工", "代碼": "2330.TW"},
        {"名稱": "Samsung Electronics", "位置": "中游", "業務": "先進邏輯製程、記憶體與晶圓代工", "代碼": "005930.KS"},
        {"名稱": "Intel", "位置": "中游", "業務": "先進節點CPU製造與晶圓代工服務", "代碼": "INTC"},
        {"名稱": "ASML", "位置": "上游", "業務": "EUV先進微影設備", "代碼": "ASML"},
        {"名稱": "東京威力科創", "位置": "上游", "業務": "先進製程蝕刻、薄膜與清洗設備", "代碼": "8035.T"},
        {"名稱": "京鼎", "位置": "上游", "業務": "先進製程設備模組與零組件", "代碼": "3413.TW"},
    ],
    "光電/材料零組件": [
        {"名稱": "大立光", "位置": "下游", "業務": "高階手機鏡頭與光學元件", "代碼": "3008.TW"},
        {"名稱": "中美晶", "位置": "上游", "業務": "半導體矽晶圓、太陽能晶圓與材料", "代碼": "5483.TWO"},
        {"名稱": "光寶科", "位置": "中游", "業務": "電源、光電元件、車用電子與伺服器電源", "代碼": "2301.TW"},
        {"名稱": "一詮", "位置": "上游", "業務": "導線架、散熱與光電零組件", "代碼": "2486.TW"},
        {"名稱": "奇鈦", "位置": "上游", "業務": "光穩定劑、紫外線吸收劑與特用化學材料", "代碼": "3430.TWO"},
    ],
    "PCB／銅箔": [
        {"名稱": "金居", "位置": "上游", "業務": "電解銅箔製造，供應高速傳輸、AI伺服器與高階PCB材料", "代碼": "8358.TWO"},
        {"名稱": "金像電", "位置": "中游", "業務": "高階多層PCB製造，應用於AI伺服器、網通與資料中心設備", "代碼": "2368.TW"},
    ],
    "ABF載板": [
        {"名稱": "味之素", "位置": "上游", "業務": "ABF絕緣膜材料與先進封裝基板關鍵材料", "代碼": "2802.T"},
        {"名稱": "欣興", "位置": "中游", "業務": "ABF載板、BT載板與高階IC封裝基板", "代碼": "3037.TW"},
        {"名稱": "南電", "位置": "中游", "業務": "ABF載板及高階IC封裝基板製造", "代碼": "8046.TW"},
        {"名稱": "景碩", "位置": "中游", "業務": "ABF載板、FC-BGA與先進IC封裝基板", "代碼": "3189.TW"},
        {"名稱": "Ibiden", "位置": "中游", "業務": "日本高階ABF載板與AI／伺服器處理器封裝基板", "代碼": "4062.T"},
        {"名稱": "Shinko Electric", "位置": "中游", "業務": "日本半導體封裝、FC-BGA及高階IC載板", "代碼": "6967.T"},
    ],
    "機器人": [
        {"名稱": "上銀", "位置": "上游", "業務": "滾珠螺桿、線性滑軌及機器人關節傳動元件", "代碼": "2049.TW"},
        {"名稱": "台灣精銳", "位置": "上游", "業務": "行星減速機與機器人腰部、關節傳動元件", "代碼": "4583.TW"},
        {"名稱": "和大", "位置": "上游", "業務": "精密齒輪與人形機器人減速機布局", "代碼": "1536.TW"},
        {"名稱": "所羅門", "位置": "中游", "業務": "AI 3D視覺軟硬體整合與機器人視覺辨識", "代碼": "2359.TW"},
        {"名稱": "達明機器人", "位置": "中游", "業務": "內建AI視覺的協作型機器人與系統整合", "代碼": "4585.TW"},
        {"名稱": "盟立", "位置": "下游", "業務": "自動化系統整合與機器人關節模組", "代碼": "2464.TW"},
        {"名稱": "鴻海", "位置": "下游", "業務": "NVIDIA Isaac平台、輪式人形與醫療協作機器人應用", "代碼": "2317.TW"},
        {"名稱": "特斯拉", "位置": "下游", "業務": "Optimus人形機器人與量產平台", "代碼": "TSLA"},
        {"名稱": "輝達", "位置": "中游", "業務": "Isaac GR00T軟體平台、Jetson Thor與機器人算力", "代碼": "NVDA"},
        {"名稱": "直覺手術", "位置": "下游", "業務": "達文西微創手術機器人系統", "代碼": "ISRG"},
        {"名稱": "泰瑞達", "位置": "下游", "業務": "Universal Robots協作機器人與工業自動化", "代碼": "TER"},
        {"名稱": "羅克韋爾自動化", "位置": "下游", "業務": "工廠自動化、工業控制及機器人整合", "代碼": "ROK"},
        {"名稱": "Symbotic", "位置": "下游", "業務": "倉儲物流AI與機器人自動化系統", "代碼": "SYM"},
        {"名稱": "宇樹科技", "位置": "下游", "業務": "中國人形與四足機器人整機", "代碼": "688836.SS"},
        {"名稱": "智元機器人", "位置": "下游", "業務": "中國人形機器人整機與具身智慧平台（未上市）", "代碼": "未上市"},
        {"名稱": "發那科", "位置": "下游", "業務": "工業機器人、CNC控制與智慧製造", "代碼": "6954.T"},
        {"名稱": "安川電機", "位置": "下游", "業務": "工業機器人、伺服驅動與運動控制", "代碼": "6506.T"},
        {"名稱": "Nabtesco", "位置": "上游", "業務": "機器人關節精密減速齒輪", "代碼": "6268.T"},
        {"名稱": "Harmonic Drive Systems", "位置": "上游", "業務": "小型與協作型機器人精密減速器", "代碼": "6324.T"},
    ],
}


@st.cache_data(ttl=600)
def fetch_stock_data(cache_version):
  items=[(category,item) for category,stocks in supply_chains.items() for item in stocks]
  tickers=list(dict.fromkeys(item["代碼"] for _,item in items if item["代碼"] != "未上市"))
  try:
    batch=yf.download(tickers,period="5d",interval="1d",group_by="ticker",auto_adjust=False,progress=False,threads=True,timeout=25)
  except Exception:
    batch=pd.DataFrame()

  results=[]
  for category,item in items:
    ticker=item["代碼"]
    if ticker == "未上市":
      results.append({
          "產業板塊":category,"股票名稱":item["名稱"],"上下游":item["位置"],"承作業務":item["業務"],"記憶體類型":item.get("記憶體類型","—"),"代碼":ticker,
          "資料狀態":"未上市／無公開行情","最新收盤價":0.0,"漲跌金額":0.0,
          "漲跌幅數值":0.0,"漲跌幅(%)":"—",
      })
      continue
    try:
      if isinstance(batch.columns,pd.MultiIndex) and ticker in batch.columns.get_level_values(0):
        close=pd.to_numeric(batch[ticker]["Close"],errors="coerce").dropna()
      elif len(tickers)==1 and "Close" in batch:
        close=pd.to_numeric(batch["Close"],errors="coerce").dropna()
      else:
        close=pd.Series(dtype=float)
      if len(close)>=2:
        close_price=float(close.iloc[-1]); prev_close=float(close.iloc[-2])
        change=close_price-prev_close; pct_change=change/prev_close*100 if prev_close else 0.0
        status="正常"
      else:
        close_price=change=pct_change=0.0; status="Yahoo暫無資料"
    except Exception:
      close_price=change=pct_change=0.0; status="讀取失敗"
    results.append({
        "產業板塊":category,"股票名稱":item["名稱"],"上下游":item["位置"],"承作業務":item["業務"],"記憶體類型":item.get("記憶體類型","—"),"代碼":ticker,
        "資料狀態":status,"最新收盤價":round(close_price,2),"漲跌金額":round(change,2),
        "漲跌幅數值":round(pct_change,2),"漲跌幅(%)":f"{round(pct_change,2)}%",
    })
  return pd.DataFrame(results)


def clean_category_label(value):
  """Hide internal maintenance tags from user-facing industry names."""
  return re.sub(r"\s*[（(](?:新增|擴充(?:板塊)?)[）)]", "", str(value)).strip()


def colored_metric(label, value, numeric_value):
  color = (
      "#ff4d4d"
      if numeric_value > 0
      else "#2eb82e" if numeric_value < 0 else "#ffffff"
  )
  arrow = "▲" if numeric_value > 0 else "▼" if numeric_value < 0 else "—"
  st.markdown(
      "<div style='background:#161b22;border:1px solid #3b4350;"
      "border-radius:10px;padding:15px 18px;min-height:92px'>"
      f"<div style='color:#f8fafc;font-size:1.05rem;font-weight:700;"
      f"margin-bottom:7px'>{label}</div>"
      f"<div style='font-size:2rem;font-weight:800;color:{color};"
      f"line-height:1.15'>{arrow} {value}</div></div>",
      unsafe_allow_html=True,
  )


def build_sector_word_report(report_df, sector_name, average_change):
  doc = Document()
  section = doc.sections[0]
  section.orientation = WD_ORIENT.LANDSCAPE
  section.page_width, section.page_height = section.page_height, section.page_width
  for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(section, attr, Inches(0.55))
  doc.styles["Normal"].font.name = "Arial"
  doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
  doc.styles["Normal"].font.size = Pt(9)

  title = doc.add_paragraph()
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run = title.add_run(f"{sector_name}｜產業板塊資料報告")
  run.bold = True; run.font.size = Pt(20); run.font.color.rgb = RGBColor(20, 55, 90)
  meta = doc.add_paragraph()
  meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
  meta.add_run(f"資料時間：{datetime.now().strftime('%Y/%m/%d %H:%M')}　板塊平均漲跌幅：{average_change:+.2f}%　來源：Yahoo Finance")

  wanted = ["股票名稱", "上下游", "承作業務", "記憶體類型", "代碼", "資料狀態", "最新收盤價", "漲跌金額", "漲跌幅(%)"]
  columns = [c for c in wanted if c in report_df.columns]
  widths = {"股票名稱":1.15,"上下游":.65,"承作業務":3.15,"記憶體類型":1.0,"代碼":.85,"資料狀態":.9,"最新收盤價":.9,"漲跌金額":.8,"漲跌幅(%)":.8}
  table = doc.add_table(rows=1, cols=len(columns)); table.style = "Table Grid"; table.autofit = False
  for idx, column in enumerate(columns):
    cell = table.rows[0].cells[idx]; cell.width = Inches(widths.get(column, 1.0)); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shading = OxmlElement("w:shd"); shading.set(qn("w:fill"), "20364F"); cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(column); r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(255,255,255)

  for _, item in report_df.iterrows():
    cells = table.add_row().cells; change = float(item.get("漲跌幅數值", 0) or 0)
    for idx, column in enumerate(columns):
      cells[idx].width = Inches(widths.get(column, 1.0)); cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
      p = cells[idx].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT if column in ("股票名稱", "承作業務") else WD_ALIGN_PARAGRAPH.CENTER
      r = p.add_run(str(item.get(column, ""))); r.font.size = Pt(8)
      if column in ("漲跌金額", "漲跌幅(%)"):
        r.bold = True; r.font.color.rgb = RGBColor(210,35,45) if change > 0 else RGBColor(20,145,75) if change < 0 else RGBColor(50,50,50)

  doc.add_paragraph("註：紅色為上漲、綠色為下跌；供應鏈位置為研究分類，並非公司官方分類。")
  output = BytesIO(); doc.save(output); output.seek(0)
  return output.getvalue()
  arrow = "▲" if numeric_value > 0 else "▼" if numeric_value < 0 else "—"
  st.markdown(
      "<div style='background:#161b22;border:1px solid #3b4350;"
      "border-radius:10px;padding:15px 18px;min-height:92px'>"
      f"<div style='color:#f8fafc;font-size:1.05rem;font-weight:700;"
      f"margin-bottom:7px'>{label}</div>"
      f"<div style='font-size:2rem;font-weight:800;color:{color};"
      f"line-height:1.15'>{arrow} {value}</div></div>",
      unsafe_allow_html=True,
  )


# 側邊欄控制面板
st.sidebar.header("🔍 控制面板")
if st.sidebar.button("🔄 重新整理即時股價"):
  st.cache_data.clear()

with st.spinner("正在從 Yahoo Finance 抓取最新跨國股價數據，請稍候..."):
  df_stocks = fetch_stock_data("20260826-optical-chain-sort-v11")
  if not df_stocks.empty:
    df_stocks["產業板塊"] = df_stocks["產業板塊"].map(clean_category_label)


# 渲染成帶有紅綠色彩的 HTML 表格函式
def render_html_table(df):
  html = "<table style='width:100%; border-collapse: collapse;'>"
  html += "<thead><tr>"
  for col in df.columns:
    html += (
        "<th style='padding:10px; background-color:#21262d; color:white;"
        f" text-align:left;'>{col}</th>"
    )
  html += "</tr></thead><tbody>"

  for _, row in df.iterrows():
    html += "<tr style='border-bottom: 1px solid #30363d;'>"
    for col in df.cols if hasattr(df, "cols") else df.columns:
      val = row[col]
      if col in ["漲跌金額", "漲跌幅(%)"]:
        try:
          num_val = float(str(row["漲跌幅(%)"]).replace("%", "").replace(",", ""))
        except (TypeError, ValueError):
          num_val = 0.0
        color = (
            "#ff4d4d"
            if num_val > 0
            else ("#2eb82e" if num_val < 0 else "white")
        )
        html += (
            f"<td style='padding:10px; color: {color}; font-weight:"
            f" bold;'>{val}</td>"
        )
      else:
        html += f"<td style='padding:10px; color: white;'>{val}</td>"
    html += "</tr>"
  html += "</tbody></table>"
  return html


# 板塊分頁：皇冠依當日正常行情的板塊平均漲跌幅動態重算。
normal_quotes = df_stocks[df_stocks["資料狀態"].eq("正常")].copy()
sector_daily = normal_quotes.groupby("產業板塊")["漲跌幅數值"].mean().sort_values(ascending=False)
best_sector = sector_daily.index[0] if not sector_daily.empty else None
best_return = float(sector_daily.iloc[0]) if not sector_daily.empty else None

page_labels = ["🌐 全部總覽"]
label_to_sector = {"🌐 全部總覽": None}
for name in supply_chains.keys():
  sector = clean_category_label(name)
  label = f"👑 {sector}" if sector == best_sector else sector
  page_labels.append(label)
  label_to_sector[label] = sector

selected_label = st.selectbox("📑 產業板塊分頁", page_labels, key="sector_page")
selected_page = label_to_sector[selected_label]
if best_sector is not None:
  st.success(f"👑 今日平均漲幅最高板塊：{best_sector}（{best_return:+.2f}%）")

if selected_page is None:
  page_df = df_stocks.copy()
  page_title = "所有產業板塊即時總覽"
else:
  page_df = df_stocks[df_stocks["產業板塊"] == selected_page].copy()
  page_title = f"{selected_page}｜板塊即時行情"

position_order = {"上游": 0, "中游": 1, "中下游": 2, "下游": 3, "跨供應鏈": 4}
page_df["_供應鏈排序"] = page_df["上下游"].map(position_order).fillna(9)
sort_columns = ["產業板塊", "_供應鏈排序"] if selected_page is None else ["_供應鏈排序"]
page_df = page_df.sort_values(sort_columns, kind="stable").drop(columns="_供應鏈排序")

st.subheader(page_title)
if page_df.empty:
  st.warning("此板塊目前沒有可顯示的公司資料。")
else:
  valid_page = page_df[page_df["資料狀態"].eq("正常")]
  avg_change = valid_page["漲跌幅數值"].mean() if not valid_page.empty else 0.0
  up_count = int((valid_page["漲跌幅數值"] > 0).sum())
  down_count = int((valid_page["漲跌幅數值"] < 0).sum())
  col1, col2, col3 = st.columns(3)
  with col1:
    average_label = "🌐 全市場平均漲跌幅" if selected_page is None else f"📈 {selected_page}平均漲跌幅"
    colored_metric(average_label, f"{avg_change:.2f}%", avg_change)
  with col2:
    st.metric(label="🔴 上漲家數", value=f"{up_count} 家")
  with col3:
    st.metric(label="🟢 下跌家數", value=f"{down_count} 家")

  report_name = selected_page or "全部產業"
  word_data = build_sector_word_report(page_df, report_name, avg_change)
  st.download_button(
      label=f"📄 下載 {report_name} Word資料表",
      data=word_data,
      file_name=f"{report_name}_產業資料_{datetime.now().strftime('%Y%m%d')}.docx",
      mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
      key=f"word_{report_name}",
  )

  display_cols = [c for c in page_df.columns if c != "漲跌幅數值"]
  display_table = page_df[display_cols].copy()

  def table_color(value):
    try:
      number = float(str(value).replace("%", "").replace(",", ""))
    except (TypeError, ValueError):
      return ""
    if number > 0:
      return "color:#d62728;font-weight:700"
    if number < 0:
      return "color:#169c46;font-weight:700"
    return ""

  styled = (
      display_table.style
      .set_properties(**{
          "color": "#111827",
          "background-color": "#ffffff",
          "font-weight": "500",
      })
      .set_table_styles([
          {"selector": "th", "props": [
              ("color", "#ffffff"),
              ("background-color", "#20252d"),
              ("font-weight", "700"),
          ]},
      ])
      .map(table_color, subset=["漲跌金額", "漲跌幅(%)"])
  )
  st.dataframe(styled, hide_index=True, width="stretch", height=650)
  st.caption("紅色＝上漲；綠色＝下跌。皇冠依當日正常行情公司的板塊平均漲跌幅重算；無行情公司不納入排名。")
